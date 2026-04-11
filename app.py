import asyncio
import io
import os
import tempfile
import threading
import uuid

from flask import Flask, jsonify, render_template, request, send_file
from PIL import Image
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import pytesseract

import crawler

app = Flask(__name__)

jobs = {}
job_lock = threading.Lock()


# ── Webpage to PDF (async crawl job) ─────────────────────────────────────────

def run_webpage_pdf_job(job_id, url, work_dir):
    import time
    start_time = time.time()

    def progress_cb(done, total, current_url=''):
        elapsed = time.time() - start_time
        rate    = done / elapsed if elapsed > 0 else 0
        eta     = int((total - done) / rate) if rate > 0 and total > done else 0
        eta_str = f'{eta // 60}m {eta % 60}s' if eta >= 60 else f'{eta}s'
        with job_lock:
            jobs[job_id].update({
                'progress': f'Crawling page {done} of {total}',
                'current_url': current_url,
                'done': done, 'total': total,
                'elapsed': int(elapsed),
                'eta': eta_str if done > 0 else 'Calculating...',
                'rate': round(rate, 2),
            })

    try:
        with job_lock:
            jobs[job_id]['status'] = 'running'
        pdf_path, domain = asyncio.run(crawler.run(url, work_dir, progress_cb))
        elapsed = int(time.time() - start_time)
        with job_lock:
            jobs[job_id].update({'status': 'done', 'file_path': pdf_path, 'domain': domain,
                                 'elapsed': elapsed, 'filename': domain.replace('.', '_') + '.pdf',
                                 'mimetype': 'application/pdf'})
    except Exception as e:
        with job_lock:
            jobs[job_id].update({'status': 'error', 'error': str(e)})


# ── Webpage to Doc ────────────────────────────────────────────────────────────

def run_webpage_doc_job(job_id, url, work_dir):
    import time
    start_time = time.time()

    def progress_cb(done, total, current_url=''):
        elapsed = time.time() - start_time
        rate    = done / elapsed if elapsed > 0 else 0
        eta     = int((total - done) / rate) if rate > 0 and total > done else 0
        eta_str = f'{eta // 60}m {eta % 60}s' if eta >= 60 else f'{eta}s'
        with job_lock:
            jobs[job_id].update({
                'progress': f'Crawling page {done} of {total}',
                'current_url': current_url,
                'done': done, 'total': total,
                'elapsed': int(elapsed),
                'eta': eta_str if done > 0 else 'Calculating...',
                'rate': round(rate, 2),
            })

    try:
        with job_lock:
            jobs[job_id]['status'] = 'running'
        visited, domain = asyncio.run(crawler.crawl_only(url, work_dir, progress_cb))

        with job_lock:
            jobs[job_id]['progress'] = 'Building Word document...'

        doc      = Document()
        doc.add_heading(domain, 0)
        for page_url, data in visited.items():
            doc.add_heading(data.get('title') or page_url, level=1)
            doc.add_paragraph(page_url).italic = True
            for item in data.get('items', []):
                if item['type'] == 'section':
                    lvl = item.get('level', 1)
                    if item.get('heading'):
                        doc.add_heading(item['heading'], level=min(lvl + 1, 4))
                    if item.get('text'):
                        doc.add_paragraph(item['text'])
                elif item['type'] == 'image':
                    path = item.get('local_path')
                    if path and os.path.exists(path):
                        try:
                            doc.add_picture(path, width=Inches(5))
                        except Exception:
                            pass
            doc.add_page_break()

        out_path = os.path.join(work_dir, domain.replace('.', '_') + '.docx')
        doc.save(out_path)
        elapsed = int(time.time() - start_time)
        with job_lock:
            jobs[job_id].update({'status': 'done', 'file_path': out_path, 'domain': domain,
                                 'elapsed': elapsed, 'filename': domain.replace('.', '_') + '.docx',
                                 'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
    except Exception as e:
        with job_lock:
            jobs[job_id].update({'status': 'error', 'error': str(e)})


# ── Image to PDF ──────────────────────────────────────────────────────────────

def images_to_pdf(images, work_dir):
    pdf = FPDF()
    pdf.set_auto_page_break(False)
    for img_bytes, name in images:
        with io.BytesIO(img_bytes) as buf:
            img = Image.open(buf).convert('RGB')
            w, h = img.size
            # A4 = 210x297mm, fit image keeping aspect ratio
            ratio  = min(190 / (w * 0.264583), 270 / (h * 0.264583), 1.0)
            w_mm   = w * 0.264583 * ratio
            h_mm   = h * 0.264583 * ratio
            tmp    = os.path.join(work_dir, f'{uuid.uuid4().hex}.jpg')
            img.save(tmp, 'JPEG', quality=85)
        pdf.add_page()
        pdf.image(tmp, x=(210 - w_mm) / 2, y=(297 - h_mm) / 2, w=w_mm, h=h_mm)
    out = os.path.join(work_dir, 'images.pdf')
    pdf.output(out)
    return out


# ── Image to Doc ──────────────────────────────────────────────────────────────

def images_to_doc(images, work_dir, ocr=True):
    doc = Document()
    doc.add_heading('WebDocs — Image Conversion', 0)
    for img_bytes, name in images:
        doc.add_heading(name, level=1)
        tmp = os.path.join(work_dir, f'{uuid.uuid4().hex}.jpg')
        with io.BytesIO(img_bytes) as buf:
            img = Image.open(buf).convert('RGB')
            img.save(tmp, 'JPEG', quality=85)
        try:
            doc.add_picture(tmp, width=Inches(5))
        except Exception:
            pass
        if ocr:
            try:
                text = pytesseract.image_to_string(Image.open(tmp)).strip()
                if text:
                    doc.add_heading('Extracted Text', level=2)
                    doc.add_paragraph(text)
            except Exception:
                pass
        doc.add_page_break()
    out = os.path.join(work_dir, 'images.docx')
    doc.save(out)
    return out


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/convert/webpage-pdf', methods=['POST'])
def convert_webpage_pdf():
    data = request.get_json()
    url  = (data or {}).get('url', '').strip()
    if not url: return jsonify({'error': 'URL is required.'}), 400
    if not url.startswith(('http://', 'https://')): url = 'https://' + url
    job_id = str(uuid.uuid4())
    work_dir = tempfile.mkdtemp()
    with job_lock:
        jobs[job_id] = {'status': 'queued', 'file_path': None, 'error': None, 'progress': 'Starting...'}
    threading.Thread(target=run_webpage_pdf_job, args=(job_id, url, work_dir), daemon=True).start()
    return jsonify({'job_id': job_id})


@app.route('/convert/webpage-doc', methods=['POST'])
def convert_webpage_doc():
    data = request.get_json()
    url  = (data or {}).get('url', '').strip()
    if not url: return jsonify({'error': 'URL is required.'}), 400
    if not url.startswith(('http://', 'https://')): url = 'https://' + url
    job_id = str(uuid.uuid4())
    work_dir = tempfile.mkdtemp()
    with job_lock:
        jobs[job_id] = {'status': 'queued', 'file_path': None, 'error': None, 'progress': 'Starting...'}
    threading.Thread(target=run_webpage_doc_job, args=(job_id, url, work_dir), daemon=True).start()
    return jsonify({'job_id': job_id})


@app.route('/convert/image-pdf', methods=['POST'])
def convert_image_pdf():
    files = request.files.getlist('images')
    if not files: return jsonify({'error': 'No images uploaded.'}), 400
    work_dir = tempfile.mkdtemp()
    images   = [(f.read(), f.filename) for f in files]
    try:
        out = images_to_pdf(images, work_dir)
        return send_file(out, as_attachment=True, download_name='webdocs_images.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/convert/image-doc', methods=['POST'])
def convert_image_doc():
    files = request.files.getlist('images')
    if not files: return jsonify({'error': 'No images uploaded.'}), 400
    work_dir = tempfile.mkdtemp()
    images   = [(f.read(), f.filename) for f in files]
    try:
        out = images_to_doc(images, work_dir)
        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return send_file(out, as_attachment=True, download_name='webdocs_images.docx', mimetype=mime)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/status/<job_id>')
def status(job_id):
    with job_lock:
        job = jobs.get(job_id)
    if not job: return jsonify({'error': 'Job not found.'}), 404
    return jsonify({
        'status':      job['status'],
        'progress':    job.get('progress', ''),
        'current_url': job.get('current_url', ''),
        'done':        job.get('done', 0),
        'total':       job.get('total', 0),
        'elapsed':     job.get('elapsed', 0),
        'eta':         job.get('eta', ''),
        'rate':        job.get('rate', 0),
        'error':       job.get('error', ''),
    })


@app.route('/download/<job_id>')
def download(job_id):
    with job_lock:
        job = jobs.get(job_id)
    if not job or job['status'] != 'done':
        return jsonify({'error': 'File not ready.'}), 404
    return send_file(job['file_path'], as_attachment=True,
                     download_name=job.get('filename', 'output'),
                     mimetype=job.get('mimetype', 'application/octet-stream'))


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
